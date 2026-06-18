from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from .styles import Fonts, Sizes, Colors, Margins, set_cell_shading


class DocxBuilder:
    def __init__(self):
        self.doc = Document()
        self._setup_default_styles()

    def _setup_default_styles(self):
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = Fonts.BODY
        font.size = Sizes.BODY
        font.color.rgb = Colors.DARK
        paragraph_format = style.paragraph_format
        paragraph_format.space_after = Pt(6)
        paragraph_format.line_spacing = 1.15

    def set_margins(self, margins=None):
        top, bottom, left, right = margins or Margins.NORMAL
        for section in self.doc.sections:
            section.top_margin = top
            section.bottom_margin = bottom
            section.left_margin = left
            section.right_margin = right

    def add_heading(self, text, level=1):
        heading = self.doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.color.rgb = Colors.PRIMARY
        return heading

    def add_paragraph(self, text, style=None, bold=False, italic=False, size=None, color=None, alignment=None, spacing_after=None):
        p = self.doc.add_paragraph(style=style)
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = size
        if color:
            run.font.color.rgb = color
        run.font.name = Fonts.BODY
        if alignment:
            p.alignment = alignment
        if spacing_after is not None:
            p.paragraph_format.space_after = spacing_after
        return p

    def add_mixed_paragraph(self, parts, alignment=None, spacing_after=None):
        p = self.doc.add_paragraph()
        for part in parts:
            run = p.add_run(part.get("text", ""))
            run.bold = part.get("bold", False)
            run.italic = part.get("italic", False)
            run.font.size = part.get("size", Sizes.BODY)
            run.font.color.rgb = part.get("color", Colors.DARK)
            run.font.name = part.get("font", Fonts.BODY)
        if alignment:
            p.alignment = alignment
        if spacing_after is not None:
            p.paragraph_format.space_after = spacing_after
        return p

    def add_table(self, headers, rows, col_widths=None, header_color=None):
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header_color = header_color or Colors.PRIMARY
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(header)
            run.bold = True
            run.font.size = Sizes.TABLE_HEADER
            run.font.color.rgb = Colors.WHITE
            run.font.name = Fonts.BODY
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            hex_color = "{:02X}{:02X}{:02X}".format(
                header_color[0], header_color[1], header_color[2]
            )
            set_cell_shading(cell, hex_color)
        for r_idx, row_data in enumerate(rows):
            for c_idx, cell_text in enumerate(row_data):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(str(cell_text))
                run.font.size = Sizes.TABLE_CELL
                run.font.name = Fonts.BODY
                run.font.color.rgb = Colors.DARK
        if col_widths:
            for row in table.rows:
                for i, width in enumerate(col_widths):
                    row.cells[i].width = width
        return table

    def add_image(self, image_path, width=None, height=None, alignment=None):
        if width and height:
            pic = self.doc.add_picture(image_path, width=width, height=height)
        elif width:
            pic = self.doc.add_picture(image_path, width=width)
        elif height:
            pic = self.doc.add_picture(image_path, height=height)
        else:
            pic = self.doc.add_picture(image_path)
        last_paragraph = self.doc.paragraphs[-1]
        if alignment:
            last_paragraph.alignment = alignment
        return pic

    def add_page_break(self):
        self.doc.add_page_break()

    def add_horizontal_rule(self):
        p = self.doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            '<w:pBdr {}>'
            '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="BFBFBF"/>'
            '</w:pBdr>'.format(nsdecls("w"))
        )
        pPr.append(pBdr)

    def add_header(self, text, alignment=None):
        section = self.doc.sections[0]
        header = section.header
        p = header.paragraphs[0]
        p.text = text
        if alignment:
            p.alignment = alignment
        for run in p.runs:
            run.font.size = Sizes.SMALL
            run.font.color.rgb = Colors.GRAY
            run.font.name = Fonts.BODY

    def add_footer(self, text, alignment=None):
        section = self.doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = text
        if alignment:
            p.alignment = alignment
        for run in p.runs:
            run.font.size = Sizes.SMALL
            run.font.color.rgb = Colors.GRAY
            run.font.name = Fonts.BODY

    def save(self, path):
        self.doc.save(path)
