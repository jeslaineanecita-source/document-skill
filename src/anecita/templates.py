from datetime import date
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .builder import DocxBuilder
from .styles import Fonts, Sizes, Colors, Margins


class ReportTemplate:
    def __init__(self, title, author=None, date_str=None):
        self.builder = DocxBuilder()
        self.title = title
        self.author = author
        self.date_str = date_str or date.today().isoformat()

    def build(self, sections):
        self.builder.set_margins(Margins.NORMAL)
        self._add_title_page()
        self.builder.add_page_break()
        for section in sections:
            self.builder.add_heading(section["heading"], level=1)
            for item in section.get("body", []):
                if item.get("type") == "paragraph":
                    self.builder.add_paragraph(
                        item["text"],
                        bold=item.get("bold"),
                        size=item.get("size"),
                        alignment=item.get("alignment"),
                    )
                elif item.get("type") == "table":
                    self.builder.add_table(
                        item["headers"],
                        item["rows"],
                        col_widths=item.get("col_widths"),
                    )
                elif item.get("type") == "image":
                    self.builder.add_image(
                        item["path"],
                        width=item.get("width"),
                        height=item.get("height"),
                        alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    )
        return self.builder

    def _add_title_page(self):
        self.builder.add_paragraph("")
        self.builder.add_paragraph("")
        self.builder.add_paragraph("")
        self.builder.add_paragraph(
            self.title,
            bold=True,
            size=Sizes.TITLE,
            color=Colors.PRIMARY,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            spacing_after=Pt(24),
        )
        if self.author:
            self.builder.add_paragraph(
                f"Author: {self.author}",
                size=Sizes.HEADING_3,
                color=Colors.GRAY,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
            )
        self.builder.add_paragraph(
            f"Date: {self.date_str}",
            size=Sizes.BODY,
            color=Colors.GRAY,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )


class InvoiceTemplate:
    def __init__(self, invoice_number, company_name, client_name, date_str=None):
        self.builder = DocxBuilder()
        self.invoice_number = invoice_number
        self.company_name = company_name
        self.client_name = client_name
        self.date_str = date_str or date.today().isoformat()

    def build(self, line_items, tax_rate=0.0, notes=None):
        self.builder.set_margins(Margins.NORMAL)
        self._add_header()
        self.builder.add_horizontal_rule()
        self._add_info_block()
        self.builder.add_heading("Invoice Details", level=2)
        subtotal = self._add_line_items_table(line_items)
        self._add_totals(subtotal, tax_rate)
        if notes:
            self.builder.add_paragraph("")
            self.builder.add_paragraph(f"Notes: {notes}", italic=True, color=Colors.GRAY)
        return self.builder

    def _add_header(self):
        self.builder.add_paragraph(
            "INVOICE",
            bold=True,
            size=Sizes.TITLE,
            color=Colors.PRIMARY,
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        )
        self.builder.add_paragraph(
            f"#{self.invoice_number}",
            size=Sizes.HEADING_2,
            color=Colors.GRAY,
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        )

    def _add_info_block(self):
        info = [
            {"text": self.company_name, "bold": True, "size": Sizes.HEADING_3},
            {"text": f"Date: {self.date_str}", "size": Sizes.BODY},
            {"text": f"Bill To: {self.client_name}", "size": Sizes.BODY},
        ]
        for item in info:
            self.builder.add_paragraph(
                item["text"],
                bold=item.get("bold", False),
                size=item.get("size", Sizes.BODY),
                color=item.get("color", Colors.DARK),
            )

    def _add_line_items_table(self, line_items):
        headers = ["#", "Description", "Qty", "Unit Price", "Total"]
        rows = []
        subtotal = 0.0
        for i, item in enumerate(line_items, 1):
            qty = item.get("qty", 1)
            unit_price = item.get("unit_price", 0)
            total = qty * unit_price
            subtotal += total
            rows.append([
                str(i),
                item["description"],
                str(qty),
                f"${unit_price:.2f}",
                f"${total:.2f}",
            ])
        self.builder.add_table(headers, rows)
        return subtotal

    def _add_totals(self, subtotal, tax_rate):
        tax = subtotal * tax_rate
        grand_total = subtotal + tax
        self.builder.add_paragraph("")
        lines = [
            ("Subtotal:", f"${subtotal:.2f}"),
            (f"Tax ({tax_rate*100:.0f}%):", f"${tax:.2f}"),
            ("Total:", f"${grand_total:.2f}"),
        ]
        for label, value in lines:
            is_bold = label == "Total:"
            self.builder.add_mixed_paragraph(
                [
                    {"text": label, "bold": is_bold, "size": Sizes.BODY, "color": Colors.DARK},
                    {"text": f"  {value}", "bold": is_bold, "size": Sizes.BODY, "color": Colors.PRIMARY},
                ],
                alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            )


class LetterTemplate:
    def __init__(self, sender, recipient, subject, date_str=None):
        self.builder = DocxBuilder()
        self.sender = sender
        self.recipient = recipient
        self.subject = subject
        self.date_str = date_str or date.today().isoformat()

    def build(self, body_paragraphs, closing="Sincerely", signature_name=None):
        self.builder.set_margins(Margins.WIDE)
        self._add_sender_block()
        self.builder.add_paragraph("")
        self.builder.add_paragraph(self.date_str)
        self.builder.add_paragraph("")
        self._add_recipient_block()
        self.builder.add_paragraph("")
        self.builder.add_paragraph(
            f"Re: {self.subject}",
            bold=True,
            size=Sizes.HEADING_3,
            color=Colors.PRIMARY,
        )
        self.builder.add_paragraph("")
        for para in body_paragraphs:
            self.builder.add_paragraph(para, size=Sizes.BODY)
            self.builder.add_paragraph("")
        self.builder.add_paragraph("")
        self.builder.add_paragraph(closing)
        self.builder.add_paragraph("")
        self.builder.add_paragraph("")
        self.builder.add_paragraph(signature_name or self.sender, bold=True)
        return self.builder

    def _add_sender_block(self):
        for line in self.sender.split("\n"):
            self.builder.add_paragraph(line, size=Sizes.BODY, color=Colors.DARK)

    def _add_recipient_block(self):
        for line in self.recipient.split("\n"):
            self.builder.add_paragraph(line, size=Sizes.BODY, color=Colors.DARK)
